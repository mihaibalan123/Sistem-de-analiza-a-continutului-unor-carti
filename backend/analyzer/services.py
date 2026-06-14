import os
import re
import json
import time
import threading
import logging
from PIL import Image
import cv2
import numpy as np
from pdf2image import convert_from_path, pdfinfo_from_path
import pytesseract
from django.db import transaction
from django.conf import settings
from .models import Carte, Autor, Personaj, Relatie
from .agents import DialogAgent
from docx import Document
from docx.shared import Pt, RGBColor
logger = logging.getLogger(__name__)
TESSERACT_PATH = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
if os.path.exists(TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
else:
    logger.warning(f"Tesseract OCR not found at '{TESSERACT_PATH}'.")
PROGRESS_LOCK = threading.Lock()
PROGRESS_STATE = {}

def get_progress(carte_id):
    with PROGRESS_LOCK:
        if carte_id in PROGRESS_STATE:
            return PROGRESS_STATE[carte_id]
    try:
        carte = Carte.objects.get(id_carte=carte_id)
        import glob
        outputs_dir = os.path.normpath(os.path.join(settings.BASE_DIR, '..', 'outputs'))
        pattern = os.path.join(outputs_dir, f'carte_{carte_id}_*.json')
        ocr_exists = len(glob.glob(pattern)) > 0
        if not ocr_exists:
            import unicodedata
            nfkd_form = unicodedata.normalize('NFKD', carte.titlu)
            ascii_str = nfkd_form.encode('ASCII', 'ignore').decode('ASCII')
            safe_chars = []
            for char in ascii_str:
                if char.isalnum() or char in ['.', '-', '_']:
                    safe_chars.append(char)
                elif char.isspace():
                    safe_chars.append('_')
            safe_name = ''.join(safe_chars)
            pattern_fallback = os.path.join(outputs_dir, f'*{safe_name}*_extras.json')
            ocr_exists = len(glob.glob(pattern_fallback)) > 0
        has_characters = Personaj.objects.filter(id_carte=carte).exists()
        if has_characters:
            return {'status': 'done', 'done': 100, 'total': 100, 'message': 'Analiză AI finalizată!'}
        elif ocr_exists:
            return {'status': 'done_ocr', 'done': carte.nr_pagini or 1, 'total': carte.nr_pagini or 1, 'message': 'OCR Finalizat. Gata pentru Analiză AI!'}
        else:
            return {'status': 'pending', 'done': 0, 'total': carte.nr_pagini or 0, 'message': 'Așteaptă procesare OCR...'}
    except Carte.DoesNotExist:
        return {'status': 'error', 'done': 0, 'total': 0, 'message': 'Cartea nu a fost găsită în sistem.'}
    except Exception as e:
        return {'status': 'error', 'done': 0, 'total': 0, 'message': f'Eroare la determinarea stării: {str(e)}'}

def update_progress(carte_id, status, done, total, message=''):
    with PROGRESS_LOCK:
        PROGRESS_STATE[carte_id] = {'status': status, 'done': done, 'total': total, 'message': message}

def update_progress_relations(carte_id, done, total, status='processing', message=''):
    with PROGRESS_LOCK:
        if carte_id not in PROGRESS_STATE:
            PROGRESS_STATE[carte_id] = {}
        state = PROGRESS_STATE[carte_id]
        state['r_done'] = done
        state['r_total'] = total
        state['r_status'] = status
        state['status'] = 'processing_ai'
        s_msg = state.get('s_message', '')
        state['r_message'] = message
        if state.get('run_summary'):
            state['message'] = f'Relații: {message} | Rezumat: {s_msg}'
        else:
            state['message'] = f'Relații: {message}'
        run_rel = state.get('run_relations', True)
        run_sum = state.get('run_summary', True)
        r_pct = done / total if total > 0 else 0
        if run_rel and run_sum:
            s_done = state.get('s_done', 0)
            s_total = state.get('s_total', 1)
            s_pct = s_done / s_total if s_total > 0 else 0
            state['done'] = int((r_pct + s_pct) / 2 * 100)
        elif run_rel:
            state['done'] = int(r_pct * 100)
        else:
            state['done'] = 100
        state['total'] = 100

def update_progress_summary(carte_id, done, total, status='processing', message=''):
    with PROGRESS_LOCK:
        if carte_id not in PROGRESS_STATE:
            PROGRESS_STATE[carte_id] = {}
        state = PROGRESS_STATE[carte_id]
        state['s_done'] = done
        state['s_total'] = total
        state['s_status'] = status
        state['status'] = 'processing_ai'
        state['s_message'] = message
        r_msg = state.get('r_message', '')
        if state.get('run_relations'):
            state['message'] = f'Relații: {r_msg} | Rezumat: {message}'
        else:
            state['message'] = f'Rezumat: {message}'
        run_rel = state.get('run_relations', True)
        run_sum = state.get('run_summary', True)
        s_pct = done / total if total > 0 else 0
        if run_rel and run_sum:
            r_done = state.get('r_done', 0)
            r_total = state.get('r_total', 1)
            r_pct = r_done / r_total if r_total > 0 else 0
            state['done'] = int((r_pct + s_pct) / 2 * 100)
        elif run_sum:
            state['done'] = int(s_pct * 100)
        else:
            state['done'] = 100
        state['total'] = 100

def check_and_finalize_ai(carte_id):
    with PROGRESS_LOCK:
        state = PROGRESS_STATE.get(carte_id, {})
        run_rel = state.get('run_relations', True)
        run_sum = state.get('run_summary', True)
        rel_done = not run_rel or state.get('r_status') == 'done'
        sum_done = not run_sum or state.get('s_status') == 'done'
        if rel_done and sum_done:
            if run_rel and run_sum:
                msg = 'Analiză AI și rezumat Word finalizate cu succes!'
            elif run_rel:
                msg = 'Analiză relații finalizată cu succes!'
            else:
                msg = 'Rezumat Word finalizat cu succes!'
            PROGRESS_STATE[carte_id] = {'status': 'done', 'done': 100, 'total': 100, 'message': msg}

def report_ai_error(carte_id, err_msg):
    with PROGRESS_LOCK:
        PROGRESS_STATE[carte_id] = {'status': 'error', 'done': 0, 'total': 0, 'message': err_msg}

def preprocesare_imagine_avansata(img_pil):
    img_cv = cv2.cvtColor(np.array(img_pil), cv2.COLOR_RGB2BGR)
    img_gri = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
    img_blur = cv2.GaussianBlur(img_gri, (3, 3), 0)
    _, img_binara = cv2.threshold(img_blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return Image.fromarray(img_binara)

def repara_paragrafe(text_brut):
    text = re.sub('-\\n\\s*', '', text_brut)
    text = re.sub('\\n\\s*([—\\-])', '\\n\\n\\1', text)
    text = re.sub('(?<!\\n)\\n(?!\\n)', ' ', text)
    text = re.sub(' {2,}', ' ', text)
    text = text.replace('\n\n', '\n')
    return text.strip()

def proceseaza_pagina_ocr(cale_pdf, numar_pagina):
    try:
        imagini = convert_from_path(cale_pdf, dpi=300, first_page=numar_pagina, last_page=numar_pagina)
        if not imagini:
            return ''
        img_pil_initial = imagini[0]
        img_pil_curata = preprocesare_imagine_avansata(img_pil_initial)
        custom_config = '--psm 6'
        text_raw = pytesseract.image_to_string(img_pil_curata, lang='ron', config=custom_config).strip()
        return repara_paragrafe(text_raw)
    except Exception as e:
        logger.error(f'Eroare la procesarea OCR a paginii {numar_pagina}: {e}')
        return ''

def start_parsing_background(pdf_path, carte_id):
    thread = threading.Thread(target=_worker_parsing_thread, args=(pdf_path, carte_id))
    thread.daemon = True
    thread.start()

def _worker_parsing_thread(pdf_path, carte_id):
    logger.info(f'Încep procesarea în fundal pentru cartea ID {carte_id}')
    try:
        info = pdfinfo_from_path(pdf_path)
        total_pagini = info['Pages']
        with transaction.atomic():
            carte = Carte.objects.get(id_carte=carte_id)
            carte.nr_pagini = total_pagini
            carte.save()
        outputs_dir = os.path.normpath(os.path.join(settings.BASE_DIR, '..', 'outputs'))
        os.makedirs(outputs_dir, exist_ok=True)
        pdf_basename = os.path.basename(pdf_path)
        pdf_name_only = os.path.splitext(pdf_basename)[0]
        json_output_path = os.path.join(outputs_dir, f'{pdf_name_only}_extras.json')
        date_carte = []
        for i in range(1, total_pagini + 1):
            if not Carte.objects.filter(id_carte=carte_id).exists():
                logger.info(f'Cartea ID {carte_id} a fost ștearsă din baza de date. Oprim parsarea.')
                break
            update_progress(carte_id, 'processing', i, total_pagini, f'Parsăm pagina {i}/{total_pagini}...')
            text_pagina = proceseaza_pagina_ocr(pdf_path, i)
            date_carte.append({'pagina_pdf': i, 'text_extras': text_pagina})
            with open(json_output_path, 'w', encoding='utf-8') as f:
                json.dump(date_carte, f, ensure_ascii=False, indent=4)
            logger.info(f'Cartea {carte_id}: Gata OCR pagina {i} / {total_pagini}')
        update_progress(carte_id, 'done_ocr', total_pagini, total_pagini, 'OCR finalizat cu succes! Gata pentru analiză AI.')
        logger.info(f'Procesare finalizată total pentru cartea ID {carte_id}')
    except Exception as e:
        logger.error(f'Eroare critică în thread-ul de procesare pentru cartea {carte_id}: {e}')
        update_progress(carte_id, 'error', 0, 0, f'Eroare la parsarea PDF: {str(e)}')

def start_ai_analysis_background(json_path, carte_id, run_relations=True, run_summary=True):
    with PROGRESS_LOCK:
        PROGRESS_STATE[carte_id] = {'status': 'processing_ai', 'done': 0, 'total': 100, 'run_relations': run_relations, 'run_summary': run_summary, 'r_done': 0, 'r_total': 1, 'r_status': 'processing' if run_relations else 'done', 'r_message': 'Inițializare...' if run_relations else 'Dezactivat', 's_done': 0, 's_total': 1, 's_status': 'processing' if run_summary else 'done', 's_message': 'Inițializare...' if run_summary else 'Dezactivat', 'message': 'Inițializare analiză AI...'}
    from .agents import DialogAgent, SummaryAgent
    dialog_agent = DialogAgent()
    summary_agent = SummaryAgent(dialog_agent=dialog_agent)
    threads = []
    if run_relations:
        thread_rel = threading.Thread(target=_worker_ai_analysis_thread, args=(json_path, carte_id, dialog_agent))
        thread_rel.daemon = True
        threads.append(thread_rel)
    if run_summary:
        thread_sum = threading.Thread(target=_worker_summary_thread, args=(json_path, carte_id, summary_agent))
        thread_sum.daemon = True
        threads.append(thread_sum)
    for t in threads:
        t.start()

def _worker_ai_analysis_thread(json_path, carte_id, agent=None):
    logger.info(f'Încep analiza de relații în fundal pentru cartea ID {carte_id} folosind fișierul {json_path}')
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            date_carte = json.load(f)
        if agent is None:
            from .agents import DialogAgent
            agent = DialogAgent()
        dimensiune_lot = 5
        loturi = [date_carte[x:x + dimensiune_lot] for x in range(0, len(date_carte), dimensiune_lot)]
        total_loturi = len(loturi)
        update_progress_relations(carte_id, 0, total_loturi, 'processing', f'Pregătire (0/{total_loturi} loturi)...')
        for idx, lot in enumerate(loturi):
            if not Carte.objects.filter(id_carte=carte_id).exists():
                logger.info(f'Cartea ID {carte_id} a fost ștearsă. Oprim analiza de relații.')
                break
            update_progress_relations(carte_id, idx, total_loturi, 'processing', f'Analizăm lotul {idx + 1}/{total_loturi}...')
            text_lot = ''
            for pag in lot:
                if pag['text_extras'].strip():
                    text_lot += f'\n\n--- PAGINA {pag['pagina_pdf']} ---\n\n{pag['text_extras']}'
            if text_lot.strip():
                try:
                    agent.proceseaza_pagina_si_salveaza(text_lot, carte_id)
                except Exception as ex_ai:
                    logger.error(f'Eroare la procesarea relațiilor lotului {idx + 1} pentru cartea {carte_id}: {ex_ai}')
            update_progress_relations(carte_id, idx + 1, total_loturi, 'processing', f'Lot {idx + 1}/{total_loturi} finalizat.')
            if idx < total_loturi - 1:
                logger.info(f'DialogAgent: Aștept 13 secunde pentru a respecta cota API Gemini...')
                for _ in range(13):
                    time.sleep(1)
                    if not Carte.objects.filter(id_carte=carte_id).exists():
                        break
        if Carte.objects.filter(id_carte=carte_id).exists():
            update_progress_relations(carte_id, total_loturi, total_loturi, 'done', 'Relații finalizate.')
            check_and_finalize_ai(carte_id)
            logger.info(f'Analiza de relații finalizată total pentru cartea ID {carte_id}')
    except Exception as e:
        logger.error(f'Eroare critică în thread-ul de relații pentru cartea {carte_id}: {e}')
        err_msg = str(e)
        if 'quota' in err_msg.lower() or '429' in err_msg:
            err_msg = 'Limita de cereri la API-ul Gemini a fost depășită (eroare 429 Quota Exceeded). Te rugăm să verifici contul tău Google AI Studio sau să încerci mai târziu.'
        elif 'api key' in err_msg.lower() or 'api_key' in err_msg.lower():
            err_msg = 'Cheia API Gemini configurată în fișierul .env este invalidă sau lipsește.'
        report_ai_error(carte_id, f'A apărut o eroare la analiza de relații: {err_msg}')

def _worker_summary_thread(json_path, carte_id, agent=None):
    logger.info(f'Încep thread-ul de rezumare AI pentru cartea ID {carte_id} folosinf fișierul {json_path}')
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            date_carte = json.load(f)
        if agent is None:
            from .agents import SummaryAgent
            agent = SummaryAgent()
        update_progress_summary(carte_id, 0, 1, 'processing', 'Se trimite textul complet către Gemini pentru rezumare generală (durată estimată: ~15-20 secunde)...')
        if not Carte.objects.filter(id_carte=carte_id).exists():
            logger.info(f'Cartea ID {carte_id} a fost ștearsă. Oprim thread-ul de rezumare.')
            return
        carte = Carte.objects.select_related('id_autor').get(id_carte=carte_id)
        autor_nume = f'{carte.id_autor.prenume or ''} {carte.id_autor.nume}'.strip()
        text_complet = ''
        for pag in date_carte:
            if pag.get('text_extras', '').strip():
                text_complet += f'\n\n--- PAGINA {pag['pagina_pdf']} ---\n\n{pag['text_extras']}'
        rezumat_ob = None
        if text_complet.strip():
            rezumat_ob = agent.extrage_rezumat(text_complet, carte.titlu, autor_nume)
        else:
            raise Exception('Nu s-a găsit text extras în fișierul OCR al cărții pentru a genera rezumatul.')
        if not rezumat_ob:
            raise Exception('Nu s-a putut obține rezumatul de la agentul AI.')
        update_progress_summary(carte_id, 0, 1, 'processing', 'Rezumat general finalizat. Se compilează documentul Word...')
        if Carte.objects.filter(id_carte=carte_id).exists():
            doc = Document()
            title_p = doc.add_paragraph()
            title_run = title_p.add_run(carte.titlu)
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(31, 78, 121)
            title_p.paragraph_format.space_after = Pt(6)
            author_p = doc.add_paragraph()
            author_text = f'Autor: {carte.id_autor.prenume or ''} {carte.id_autor.nume}'.strip()
            if carte.an_aparitie:
                author_text += f' | An apariție: {carte.an_aparitie}'
            author_run = author_p.add_run(author_text)
            author_run.font.name = 'Times New Roman'
            author_run.font.size = Pt(14)
            author_run.font.italic = True
            author_run.font.color.rgb = RGBColor(100, 100, 100)
            author_p.paragraph_format.space_after = Pt(24)
            doc.add_paragraph().add_run('Rezumat de Analiză Literară Inteligentă').font.bold = True
            doc.add_paragraph('Acest document conține rezumatul structurat generat automat pe baza analizei textului extras prin OCR.').paragraph_format.space_after = Pt(24)
            doc.add_page_break()
            h_gen = doc.add_heading(level=2)
            h_gen_run = h_gen.add_run('Rezumat General')
            h_gen_run.font.name = 'Times New Roman'
            h_gen_run.font.color.rgb = RGBColor(31, 78, 121)
            h_gen.paragraph_format.space_before = Pt(12)
            h_gen.paragraph_format.space_after = Pt(6)
            for paragraph_text in rezumat_ob.rezumat_general.split('\n'):
                p_text = paragraph_text.strip()
                if p_text:
                    p = doc.add_paragraph()
                    p_run = p.add_run(p_text)
                    p_run.font.name = 'Times New Roman'
                    p_run.font.size = Pt(12)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
            h_teme = doc.add_heading(level=2)
            h_teme_run = h_teme.add_run('Teme Principale')
            h_teme_run.font.name = 'Times New Roman'
            h_teme_run.font.color.rgb = RGBColor(31, 78, 121)
            h_teme.paragraph_format.space_before = Pt(18)
            h_teme.paragraph_format.space_after = Pt(6)
            for tema in rezumat_ob.teme_principale:
                if tema.strip():
                    p_tema = doc.add_paragraph(style='List Bullet')
                    p_tema_run = p_tema.add_run(tema.strip())
                    p_tema_run.font.name = 'Times New Roman'
                    p_tema_run.font.size = Pt(12)
                    p_tema.paragraph_format.space_after = Pt(4)
            doc.add_page_break()
            h_analiza = doc.add_heading(level=2)
            h_analiza_run = h_analiza.add_run('Analiză Detaliată pe Secțiuni')
            h_analiza_run.font.name = 'Times New Roman'
            h_analiza_run.font.color.rgb = RGBColor(31, 78, 121)
            h_analiza.paragraph_format.space_before = Pt(12)
            h_analiza.paragraph_format.space_after = Pt(12)
            for sec in rezumat_ob.sectiuni:
                h_sec = doc.add_heading(level=3)
                h_sec_run = h_sec.add_run(sec.titlu_sectiune)
                h_sec_run.font.name = 'Times New Roman'
                h_sec_run.font.color.rgb = RGBColor(31, 78, 121)
                h_sec.paragraph_format.space_before = Pt(12)
                h_sec.paragraph_format.space_after = Pt(6)
                if sec.idei_principale:
                    p_intro = doc.add_paragraph()
                    p_intro_run = p_intro.add_run('Idei principale:')
                    p_intro_run.font.name = 'Times New Roman'
                    p_intro_run.font.size = Pt(11)
                    p_intro_run.font.bold = True
                    p_intro.paragraph_format.space_after = Pt(4)
                    for idee in sec.idei_principale:
                        if idee.strip():
                            p_idee = doc.add_paragraph(style='List Bullet')
                            p_idee_run = p_idee.add_run(idee.strip())
                            p_idee_run.font.name = 'Times New Roman'
                            p_idee_run.font.size = Pt(11)
                            p_idee.paragraph_format.space_after = Pt(2)
                if sec.rezumat_sectiune.strip():
                    for paragraph_text in sec.rezumat_sectiune.split('\n'):
                        p_text = paragraph_text.strip()
                        if p_text:
                            p_sec = doc.add_paragraph()
                            p_sec_run = p_sec.add_run(p_text)
                            p_sec_run.font.name = 'Times New Roman'
                            p_sec_run.font.size = Pt(12)
                            p_sec.paragraph_format.space_before = Pt(6)
                            p_sec.paragraph_format.space_after = Pt(12)
            media_dir = os.path.join(settings.MEDIA_ROOT, 'summaries')
            os.makedirs(media_dir, exist_ok=True)
            docx_path = os.path.join(media_dir, f'rezumat_carte_{carte_id}.docx')
            doc.save(docx_path)
            logger.info(f'Fișierul Word a fost salvat cu succes la {docx_path}')
            update_progress_summary(carte_id, 1, 1, 'done', 'Rezumat Word finalizat.')
            check_and_finalize_ai(carte_id)
    except Exception as e:
        logger.error(f'Eroare critică în thread-ul de rezumare pentru cartea {carte_id}: {e}')
        err_msg = str(e)
        if 'quota' in err_msg.lower() or '429' in err_msg:
            err_msg = 'Limita de cereri la API-ul Gemini a fost depășită (eroare 429 Quota Exceeded). Te rugăm să reîncerci mai târziu.'
        elif 'api key' in err_msg.lower() or 'api_key' in err_msg.lower():
            err_msg = 'Cheia API Gemini configurată în fișierul .env este invalidă sau lipsește.'
        report_ai_error(carte_id, f'A apărut o eroare la generarea rezumatului: {err_msg}')